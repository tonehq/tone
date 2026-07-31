from __future__ import annotations

import io
import os
import tempfile
import time
from abc import ABC, abstractmethod
from typing import List, Optional, Tuple

from docling.datamodel.base_models import InputFormat
from docling.datamodel.layout_model_specs import DOCLING_LAYOUT_EGRET_MEDIUM
from docling.datamodel.pipeline_options import (
    AcceleratorDevice,
    AcceleratorOptions,
    LayoutOptions,
    PdfPipelineOptions,
)
from docling.document_converter import DocumentConverter, PdfFormatOption
from docx import Document as _Docx
from loguru import logger
from PyPDF2 import PdfReader as _PdfReader

from core.services.rag.types import Document


class DocumentReader(ABC):
    @abstractmethod
    def supports(self, content_type: str) -> bool:
        ...

    @abstractmethod
    def read(self, file_bytes: bytes, content_type: str, page_range: Optional[Tuple[int, int]] = None) -> Document:
        ...

    def read_path(self, file_path: str, content_type: str, page_range: Optional[Tuple[int, int]] = None) -> Document:
        with open(file_path, "rb") as f:
            return self.read(f.read(), content_type, page_range)


class DoclingReader(DocumentReader):
    _EXT = {
        "application/pdf": ".pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
        "application/msword": ".doc",
        "application/vnd.openxmlformats-officedocument.presentationml.presentation": ".pptx",
        "application/vnd.ms-powerpoint": ".ppt",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": ".xlsx",
        "application/vnd.ms-excel": ".xls",
        "text/html": ".html",
        "text/markdown": ".md",
        "text/plain": ".md",
        "text/csv": ".csv",
        "image/png": ".png",
        "image/jpeg": ".jpg",
        "image/tiff": ".tiff",
    }

    def __init__(self, page_range: Optional[Tuple[int, int]] = None, ocr: bool = True, tables: bool = True):
        self._converter = None
        self._page_range = page_range
        self._ocr = ocr
        self._tables = tables

    def _get_converter(self):
        if self._converter is None:
            num_threads = int(os.getenv("DOCLING_NUM_THREADS", "4"))
            options = PdfPipelineOptions(
                do_ocr=self._ocr,
                do_table_structure=self._tables,
                accelerator_options=AcceleratorOptions(num_threads=num_threads, device=AcceleratorDevice.CPU),
                layout_options=LayoutOptions(model_spec=DOCLING_LAYOUT_EGRET_MEDIUM),
            )
            self._converter = DocumentConverter(
                format_options={InputFormat.PDF: PdfFormatOption(pipeline_options=options)}
            )
        return self._converter

    def supports(self, content_type: str) -> bool:
        return content_type in self._EXT

    def read(self, file_bytes: bytes, content_type: str, page_range: Optional[Tuple[int, int]] = None) -> Document:
        pr = page_range or self._page_range
        ext = self._EXT.get(content_type, "")
        with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name
        kwargs = {"page_range": pr} if pr else {}
        logger.info(
            "[parse:docling] parsing {} ({:.2f} MB), page_range={}, ocr={}, tables={} ...",
            content_type, len(file_bytes) / 1024 / 1024, pr or "all", self._ocr, self._tables,
        )
        start = time.monotonic()
        try:
            result = self._get_converter().convert(tmp_path, **kwargs)
            text = result.document.export_to_markdown()
        except Exception:
            logger.exception(
                "[parse:docling] convert failed content_type={} bytes={} page_range={}",
                content_type, len(file_bytes), pr or "all",
            )
            raise
        finally:
            try:
                os.remove(tmp_path)
            except OSError:
                pass
        logger.info(
            "[parse:docling] parsed {} -> {} chars in {:.1f}s (page_range={})",
            content_type, len(text), time.monotonic() - start, pr or "all",
        )
        return Document(text=text, native=result.document, metadata={"parser": "docling", "page_range": pr})

    def read_path(self, file_path: str, content_type: str, page_range: Optional[Tuple[int, int]] = None) -> Document:
        pr = page_range or self._page_range
        kwargs = {"page_range": pr} if pr else {}
        logger.info(
            "[parse:docling] parsing path {} ({}), page_range={}, ocr={}, tables={} ...",
            file_path, content_type, pr or "all", self._ocr, self._tables,
        )
        start = time.monotonic()
        try:
            result = self._get_converter().convert(file_path, **kwargs)
            text = result.document.export_to_markdown()
        except Exception:
            logger.exception(
                "[parse:docling] convert failed path={} content_type={} page_range={}",
                file_path, content_type, pr or "all",
            )
            raise
        logger.info(
            "[parse:docling] parsed {} -> {} chars in {:.1f}s (page_range={})",
            content_type, len(text), time.monotonic() - start, pr or "all",
        )
        return Document(text=text, native=result.document, metadata={"parser": "docling", "page_range": pr})


class PdfReader(DocumentReader):
    def supports(self, content_type: str) -> bool:
        return content_type == "application/pdf"

    def read(self, file_bytes: bytes, content_type: str, page_range: Optional[Tuple[int, int]] = None) -> Document:
        try:
            reader = _PdfReader(io.BytesIO(file_bytes))
            pages = [p.extract_text() for p in reader.pages]
        except Exception:
            logger.exception(
                "[parse:pypdf] extract failed bytes={} content_type={}",
                len(file_bytes), content_type,
            )
            raise
        result = "\n\n".join(t for t in pages if t)
        logger.info("[parse:pypdf] extracted {} chars from PDF ({} pages)", len(result), len(reader.pages))
        return Document(text=result, metadata={"parser": "pypdf2"})


class DocxReader(DocumentReader):
    _TYPES = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    )

    def supports(self, content_type: str) -> bool:
        return content_type in self._TYPES

    def read(self, file_bytes: bytes, content_type: str, page_range: Optional[Tuple[int, int]] = None) -> Document:
        try:
            doc = _Docx(io.BytesIO(file_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        except Exception:
            logger.exception(
                "[parse:docx] extract failed bytes={} content_type={}",
                len(file_bytes), content_type,
            )
            raise
        result = "\n\n".join(paragraphs)
        logger.info("[parse:docx] extracted {} chars from DOCX ({} paragraphs)", len(result), len(paragraphs))
        return Document(text=result, metadata={"parser": "python-docx"})


class TextReader(DocumentReader):
    def supports(self, content_type: str) -> bool:
        return content_type.startswith("text/") or content_type in ("application/csv", "application/json")

    def read(self, file_bytes: bytes, content_type: str, page_range: Optional[Tuple[int, int]] = None) -> Document:
        return Document(text=file_bytes.decode("utf-8", errors="replace"), metadata={"parser": "text"})


class CompositeReader(DocumentReader):
    def __init__(self, readers: List[DocumentReader] = None):
        self._readers = readers or [DoclingReader(), PdfReader(), DocxReader(), TextReader()]

    def supports(self, content_type: str) -> bool:
        return any(r.supports(content_type) for r in self._readers)

    def read(self, file_bytes: bytes, content_type: str, page_range: Optional[Tuple[int, int]] = None) -> Document:
        for reader in self._readers:
            if reader.supports(content_type):
                return reader.read(file_bytes, content_type, page_range)
        raise ValueError(f"Unsupported content type for text extraction: {content_type}")

    def read_path(self, file_path: str, content_type: str, page_range: Optional[Tuple[int, int]] = None) -> Document:
        for reader in self._readers:
            if reader.supports(content_type):
                return reader.read_path(file_path, content_type, page_range)
        raise ValueError(f"Unsupported content type for text extraction: {content_type}")
